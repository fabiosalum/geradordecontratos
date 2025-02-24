import tkinter as tk
from tkinter import filedialog, messagebox
from tkinter.ttk import Progressbar
from docx import Document
from datetime import datetime
import pandas as pd
import os
import locale
from PIL import Image, ImageTk
import threading
import win32com.client
import time

class ContractGenerator:
    def __init__(self, root):
        self.root = root
        self.root.title("Gerador de Contratos Personalizado")
        self.root.geometry("550x750")
        
        # Configurar o idioma para português (mês por extenso)
        try:
            locale.setlocale(locale.LC_TIME, "pt_BR.utf8")
        except locale.Error:
            messagebox.showwarning("Aviso", "Idioma português não configurado no sistema. Usando padrão.")
        
        self.setup_gui()
        
    def setup_gui(self):
        # Frame do logo
        self.setup_logo()
        
        # Entradas principais
        self.entrada_base_dados = self.criar_linha("Base de Dados (Excel):", self.selecionar_base_dados)
        self.entrada_contrato = self.criar_linha("Modelo de Contrato (Word):", self.selecionar_contrato)
        self.entrada_pasta_destino = self.criar_linha("Pasta de Destino:", self.selecionar_pasta_destino)
        
        # Frame de variáveis
        self.setup_variaveis_frame()
        
        # Barra de progresso
        self.frame_progresso = tk.Frame(self.root)
        self.frame_progresso.pack(pady=10, padx=20, fill="x")
        self.progresso = Progressbar(self.frame_progresso, length=300, mode='determinate')
        self.progresso.pack(fill="x")
        self.label_progresso = tk.Label(self.frame_progresso, text="")
        self.label_progresso.pack()
        
        # Botões
        self.setup_botoes()
    
    def setup_logo(self):
        frame_logo = tk.Frame(self.root)
        frame_logo.pack(pady=10)
        
        try:
            imagem_path = "logo.png"
            if os.path.exists(imagem_path):
                imagem = Image.open(imagem_path)
                imagem = imagem.resize((235, 150))
                self.logo = ImageTk.PhotoImage(imagem)  # Mantém referência
                label_logo = tk.Label(frame_logo, image=self.logo)
                label_logo.pack()
            else:
                tk.Label(frame_logo, text="Logotipo não encontrado", fg="red").pack()
        except Exception as e:
            tk.Label(frame_logo, text=f"Erro ao carregar logotipo: {e}", fg="red").pack()
    
    def criar_linha(self, titulo, comando):
        frame = tk.Frame(self.root)
        frame.pack(pady=10, fill="x", padx=20)
        tk.Label(frame, text=titulo).pack(anchor="w")
        entrada = tk.Entry(frame, width=50)
        entrada.pack(side="left", expand=True, fill="x")
        tk.Button(frame, text="Selecionar", command=comando).pack(side="right")
        return entrada
    
    def setup_variaveis_frame(self):
        frame_variaveis = tk.Frame(self.root)
        frame_variaveis.pack(pady=10, padx=20, fill="x")
        tk.Label(frame_variaveis, text="Variáveis Personalizadas", font=("Arial", 12, "bold")).pack()
        
        self.entradas_variaveis = []
        for i in range(10):
            frame = tk.Frame(frame_variaveis)
            frame.pack(pady=2, fill="x")
            tk.Label(frame, text=f"Variável {i+1}:").pack(side="left")
            var_nome = tk.Entry(frame, width=15)
            var_nome.pack(side="left", padx=5)
            col_nome = tk.Entry(frame, width=25)
            col_nome.pack(side="left", padx=5)
            self.entradas_variaveis.append((var_nome, col_nome))
        
        legenda = """Legenda:
- Esquerda: Nome da variável no contrato (@exemplo)
- Direita: Nome da coluna correspondente no Excel
- Variáveis automáticas de data:
  @dd (dia), @mm (mês) e @ano (ano)"""
        
        tk.Label(frame_variaveis, text=legenda, anchor="w", justify="left", fg="red").pack(anchor="w", pady=10)
    
    def setup_botoes(self):
        frame_botoes = tk.Frame(self.root)
        frame_botoes.pack(pady=20)
        tk.Button(
            frame_botoes,
            text="Gerar Contratos em DOCX",
            command=lambda: self.iniciar_geracao("DOCX"),
            bg="green",
            fg="white"
        ).pack(side="left", padx=10)
        
        tk.Button(
            frame_botoes,
            text="Gerar Contratos em PDF",
            command=lambda: self.iniciar_geracao("PDF"),
            bg="blue",
            fg="white"
        ).pack(side="left", padx=10)
    
    def selecionar_base_dados(self):
        arquivo = filedialog.askopenfilename(filetypes=[("Planilhas Excel", "*.xlsx")])
        if arquivo:
            self.entrada_base_dados.delete(0, tk.END)
            self.entrada_base_dados.insert(0, arquivo)
    
    def selecionar_contrato(self):
        arquivo = filedialog.askopenfilename(filetypes=[("Documentos Word", "*.docx")])
        if arquivo:
            self.entrada_contrato.delete(0, tk.END)
            self.entrada_contrato.insert(0, arquivo)
    
    def selecionar_pasta_destino(self):
        pasta = filedialog.askdirectory()
        if pasta:
            self.entrada_pasta_destino.delete(0, tk.END)
            self.entrada_pasta_destino.insert(0, pasta)
    
    def validar_dados(self, tabela, referencias):
        if tabela.empty:
            return False, "A planilha está vazia!"
        
        for var, col in referencias.items():
            if col not in tabela.columns:
                return False, f"Coluna '{col}' não encontrada na planilha"
        return True, ""
    
    def substituir_texto_paragrafo(self, paragrafo, referencias):
        for run in paragrafo.runs:
            for codigo, valor in referencias.items():
                if codigo in run.text:
                    run.text = run.text.replace(codigo, str(valor))
    
    def atualizar_progresso(self, valor, mensagem):
        self.progresso["value"] = valor
        self.label_progresso["text"] = mensagem
        self.root.update_idletasks()
    
    def iniciar_geracao(self, formato):
        # Desabilitar botões durante a geração
        for widget in self.root.winfo_children():
            if isinstance(widget, tk.Button):
                widget["state"] = "disabled"
        
        # Iniciar thread de geração
        thread = threading.Thread(target=lambda: self.gerar_contratos(formato))
        thread.start()
    
    def converter_para_pdf(self, caminho_docx):
        try:
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            
            doc = word.Documents.Open(caminho_docx)
            caminho_pdf = caminho_docx.replace(".docx", ".pdf")
            
            # PDF quality settings
            doc.ExportAsFixedFormat(
                OutputFileName=caminho_pdf,
                ExportFormat=17,  # wdExportFormatPDF
                OptimizeFor=0,    # wdExportOptimizeForPrint
                CreateBookmarks=0, # wdExportCreateNoBookmarks
                DocStructureTags=True,
                BitmapMissingFonts=True,
                UseISO19005_1=False
            )
            
            doc.Close(False)
            word.Quit()
            return True, ""
        except Exception as e:
            return False, str(e)
        finally:
            try:
                doc.Close(False)
                word.Quit()
            except:
                pass
    
    def gerar_contratos(self, formato):
        try:
            # Validar campos obrigatórios
            base_dados = self.entrada_base_dados.get()
            contrato_modelo = self.entrada_contrato.get()
            pasta_destino = self.entrada_pasta_destino.get()
            
            if not all([base_dados, contrato_modelo, pasta_destino]):
                raise ValueError("Todos os campos devem ser preenchidos!")
            
            # Carregar dados
            self.atualizar_progresso(10, "Carregando dados...")
            tabela = pd.read_excel(base_dados)
            
            # Coletar referências de variáveis
            referencias = {}
            for var_nome, col_nome in self.entradas_variaveis:
                var = var_nome.get().strip()
                col = col_nome.get().strip()
                if var and col:
                    referencias[var] = col
            
            # Validar dados
            valido, mensagem = self.validar_dados(tabela, referencias)
            if not valido:
                raise ValueError(mensagem)
            
            total_registros = len(tabela.index)
            for idx, linha in enumerate(tabela.index):
                # Calcular progresso
                progresso = int((idx + 1) / total_registros * 100)
                self.atualizar_progresso(progresso, f"Processando contrato {idx + 1} de {total_registros}")
                
                documento = Document(contrato_modelo)
                valores_substituir = {}
                
                # Preparar valores para substituição
                for var, col in referencias.items():
                    valores_substituir[var] = str(tabela.loc[linha, col]).strip()
                
                # Adicionar variáveis de data
                valores_substituir.update({
                    "@dd": str(datetime.now().day),
                    "@mm": datetime.now().strftime("%B").capitalize(),
                    "@ano": str(datetime.now().year)
                })
                
                # Criar pasta individual
                nome_pasta = valores_substituir.get("@nome", f"Contrato_{idx+1}").replace("/", "_")
                pasta_pessoa = os.path.join(pasta_destino, nome_pasta)
                os.makedirs(pasta_pessoa, exist_ok=True)
                
                # Substituir textos
                for paragrafo in documento.paragraphs:
                    self.substituir_texto_paragrafo(paragrafo, valores_substituir)
                
                for tabela in documento.tables:
                    for linha in tabela.rows:
                        for celula in linha.cells:
                            for paragrafo in celula.paragraphs:
                                self.substituir_texto_paragrafo(paragrafo, valores_substituir)
                
                # Salvar documento
                caminho_docx = os.path.join(pasta_pessoa, f"Contrato - {nome_pasta}.docx")
                documento.save(caminho_docx)
                
                # Converter para PDF se necessário
                if formato == "PDF":
                    self.atualizar_progresso(progresso, f"Convertendo para PDF: {idx + 1} de {total_registros}")
                    sucesso, erro = self.converter_para_pdf(os.path.abspath(caminho_docx))
                    if not sucesso:
                        messagebox.showwarning("Aviso", f"Erro ao converter para PDF: {erro}")
                    # Aguarda um momento para garantir que o Word fechou corretamente
                    time.sleep(1)
            
            self.atualizar_progresso(100, "Concluído!")
            messagebox.showinfo("Sucesso", f"Contratos gerados com sucesso em {formato}!")
            
        except Exception as e:
            messagebox.showerror("Erro", str(e))
        finally:
            # Reabilitar botões
            for widget in self.root.winfo_children():
                if isinstance(widget, tk.Button):
                    widget["state"] = "normal"
            
            # Resetar barra de progresso
            self.atualizar_progresso(0, "")

if __name__ == "__main__":
    root = tk.Tk()
    app = ContractGenerator(root)
    root.mainloop()